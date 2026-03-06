"""Generate .docx digest report for printing."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import DigestResult

logger = logging.getLogger(__name__)


def render_digest_docx(digest: DigestResult, output_path: str, min_relevance: int = 0) -> str:
    """Render digest result to a .docx file. Returns the output path."""
    doc = Document()

    # Page margins for printing
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Title
    title = doc.add_heading("Podcast Digest", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date and stats
    stats = doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stats.add_run(f"{digest.date}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    stats.add_run("\n")
    run2 = stats.add_run(
        f"{digest.total_channels_checked} canais verificados  |  "
        f"{digest.total_new_episodes} novos episodios"
    )
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    # --- Person episodes ---
    has_person_content = digest.person_episodes or getattr(digest, "person_no_transcript", [])
    if has_person_content:
        _add_section_heading(doc, "Pessoas Rastreadas", RGBColor(0x80, 0x5A, 0xD5))

        if digest.person_names:
            p = doc.add_paragraph()
            run = p.add_run(", ".join(digest.person_names))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x5A, 0xD5)

        for ep in sorted(digest.person_episodes, key=lambda e: e.relevance_score, reverse=True):
            _add_episode_card(doc, ep, RGBColor(0x80, 0x5A, 0xD5))

        person_no_transcript = getattr(digest, "person_no_transcript", [])
        if person_no_transcript:
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            run = p.add_run("Videos encontrados (sem transcricao disponivel):")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x5A, 0xD5)
            run.italic = True
            for v in person_no_transcript:
                dur_str = f" ({v.duration_seconds // 60}min)" if v.duration_seconds else ""
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{v.title}{dur_str}")
                run.font.size = Pt(10)
                p.add_run(f"\n{v.url}").font.size = Pt(8)

    # --- Filter by minimum relevance ---
    included = [e for e in digest.episodes if e.relevance_score >= min_relevance]
    excluded_count = len(digest.episodes) - len(included)

    high = [e for e in included if e.relevance_score >= 7]
    medium = [e for e in included if 4 <= e.relevance_score < 7]
    low = [e for e in included if e.relevance_score < 4]

    if high:
        _add_section_heading(doc, "Alta Relevancia", RGBColor(0x48, 0xBB, 0x78))
        for ep in sorted(high, key=lambda e: e.relevance_score, reverse=True):
            _add_episode_card(doc, ep, RGBColor(0x48, 0xBB, 0x78))

    # --- Cross-synthesis ---
    if digest.cross_synthesis and digest.cross_synthesis.themes:
        _add_section_heading(doc, "Temas Transversais", RGBColor(0x2B, 0x6C, 0xB0))
        for theme in digest.cross_synthesis.themes:
            p = doc.add_paragraph()
            run = p.add_run(theme.get("theme", ""))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
            p.add_run("\n")
            run2 = p.add_run(theme.get("summary", ""))
            run2.font.size = Pt(10)

    # --- Medium relevance ---
    if medium:
        _add_section_heading(doc, "Relevancia Moderada", RGBColor(0xB7, 0x79, 0x1F))
        for ep in sorted(medium, key=lambda e: e.relevance_score, reverse=True):
            _add_episode_card(doc, ep, RGBColor(0xB7, 0x79, 0x1F), truncate_summary=200)

    # --- Low relevance ---
    if low:
        _add_section_heading(doc, "Baixa Relevancia", RGBColor(0xA0, 0xAE, 0xC0))
        for ep in sorted(low, key=lambda e: e.relevance_score, reverse=True):
            p = doc.add_paragraph()
            run = p.add_run(f"[{ep.relevance_score}/10] ")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
            run2 = p.add_run(ep.title)
            run2.font.size = Pt(10)
            p.add_run(f"  ({ep.channel_name})")
            p.add_run("\n")
            link_run = p.add_run(ep.url)
            link_run.font.size = Pt(8)
            link_run.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)

    # --- No transcript ---
    if digest.no_transcript:
        _add_section_heading(doc, "Sem transcricao disponivel", RGBColor(0x92, 0x40, 0x0E))
        for v in digest.no_transcript:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(v.title).font.size = Pt(10)
            p.add_run(f"\n{v.url}").font.size = Pt(8)

    # --- Footer ---
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = [f"Gerado por podcast-digest"]
    if excluded_count > 0:
        parts.append(f"{excluded_count} episodios omitidos (score < {min_relevance})")
    run = footer.add_run("  |  ".join(parts))
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

    # Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"Digest saved to {output_path}")
    return output_path


def _add_section_heading(doc: Document, text: str, color: RGBColor):
    """Add a colored section heading with a bottom border."""
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = color


def _add_episode_card(doc, ep, color: RGBColor, truncate_summary: int = 0):
    """Add an episode card to the document."""
    p = doc.add_paragraph()
    p.space_before = Pt(8)

    # Score badge + title
    badge = p.add_run(f"[{ep.relevance_score}/10]  ")
    badge.bold = True
    badge.font.size = Pt(11)
    badge.font.color.rgb = color

    title_run = p.add_run(ep.title)
    title_run.bold = True
    title_run.font.size = Pt(11)

    # Channel + duration
    p.add_run("\n")
    duration_str = f"  |  {ep.duration_seconds // 60}min" if ep.duration_seconds else ""
    meta = p.add_run(f"{ep.channel_name}{duration_str}")
    meta.font.size = Pt(9)
    meta.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    # URL
    p.add_run("\n")
    link = p.add_run(ep.url)
    link.font.size = Pt(8)
    link.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)

    # Summary
    summary_text = ep.summary
    if truncate_summary and len(summary_text) > truncate_summary:
        summary_text = summary_text[:truncate_summary] + "..."

    summary_p = doc.add_paragraph()
    run = summary_p.add_run(summary_text)
    run.font.size = Pt(10)

    # Topics
    if ep.key_topics:
        topics_p = doc.add_paragraph()
        run = topics_p.add_run("Topicos: " + ", ".join(ep.key_topics))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    # Relevance reason
    if hasattr(ep, "relevance_reason") and ep.relevance_reason:
        reason_p = doc.add_paragraph()
        run = reason_p.add_run(ep.relevance_reason)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        run.italic = True
